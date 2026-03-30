"""
Document generator — produces format-preserving optimized resumes.

Strategy:
  1. If original is DOCX: load with python-docx, replace text paragraph-by-paragraph
     while keeping ALL Run-level formatting (font, size, bold, italic, color).
     Export as DOCX and convert to PDF.

  2. If original is PDF: convert to DOCX first (pdf2docx), then follow step 1.

  3. Fallback: generate a clean, professional DOCX from scratch using the
     AI-optimized text, then convert to PDF.

PDF conversion priority:
  1. docx2pdf (uses Microsoft Word on Mac if installed)
  2. LibreOffice soffice command (if installed)
  3. ReportLab clean-render fallback (no text overlap)
"""
import io
import os
import re
import uuid
import logging
import shutil
import subprocess
from copy import deepcopy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


# ─── PDF conversion helpers ────────────────────────────────────────────────────

def _convert_docx_to_pdf_via_docx2pdf(docx_path: str, pdf_path: str) -> bool:
    """Use docx2pdf (Microsoft Word on Mac) to convert."""
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0
    except Exception as e:
        logger.warning(f"docx2pdf failed: {e}")
        return False


def _convert_docx_to_pdf_via_libreoffice(docx_path: str, out_dir: str) -> Optional[str]:
    """Use LibreOffice headless to convert DOCX → PDF."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    try:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            capture_output=True, text=True, timeout=60,
        )
        # LibreOffice places the PDF next to the DOCX with same stem
        stem = Path(docx_path).stem
        expected = os.path.join(out_dir, f"{stem}.pdf")
        if os.path.exists(expected):
            return expected
        logger.warning(f"LibreOffice output: {result.stdout} {result.stderr}")
        return None
    except Exception as e:
        logger.warning(f"LibreOffice conversion failed: {e}")
        return None


def _convert_docx_to_pdf_reportlab_fallback(docx_path: str, pdf_path: str) -> bool:
    """
    Last-resort: read the DOCX as text and render a clean PDF with ReportLab.
    No overlapping text — properly handles wrapping and page breaks.
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        doc_in = Document(docx_path)

        # Build styles
        name_style = ParagraphStyle("Name", fontSize=20, fontName="Helvetica-Bold",
                                    alignment=TA_CENTER, spaceAfter=4)
        contact_style = ParagraphStyle("Contact", fontSize=9, fontName="Helvetica",
                                       alignment=TA_CENTER, textColor=colors.HexColor("#555555"), spaceAfter=6)
        heading_style = ParagraphStyle("Heading", fontSize=11, fontName="Helvetica-Bold",
                                       textColor=colors.HexColor("#1e40af"), spaceBefore=10,
                                       spaceAfter=3, borderPadding=0)
        body_style = ParagraphStyle("Body", fontSize=9.5, fontName="Helvetica",
                                    leading=14, spaceAfter=2, leftIndent=0)
        bullet_style = ParagraphStyle("Bullet", fontSize=9.5, fontName="Helvetica",
                                      leading=14, spaceAfter=2, leftIndent=16, firstLineIndent=-10)

        SECTION_KEYWORDS = {
            "experience", "education", "skills", "projects", "summary",
            "objective", "certifications", "awards", "work experience",
            "professional experience", "technical skills", "volunteer",
        }

        story = []
        is_first_para = True
        is_second_para = False  # contact line
        section_count = 0

        for i, para in enumerate(doc_in.paragraphs):
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 4))
                continue

            lower = text.lower().rstrip(":")
            is_section = lower in SECTION_KEYWORDS or (
                len(text) < 50 and text.isupper()
            )

            if is_first_para:
                # Name line
                story.append(Paragraph(text, name_style))
                story.append(HRFlowable(width="100%", thickness=1.5,
                                        color=colors.HexColor("#1e40af"), spaceAfter=2))
                is_first_para = False
                is_second_para = True
                continue

            if is_second_para and "@" in text or (is_second_para and "|" in text):
                story.append(Paragraph(text, contact_style))
                is_second_para = False
                section_count += 1
                continue

            is_second_para = False

            if is_section:
                section_count += 1
                if section_count > 1:
                    story.append(Spacer(1, 4))
                story.append(Paragraph(text, heading_style))
                story.append(HRFlowable(width="100%", thickness=0.5,
                                        color=colors.HexColor("#e5e7eb"), spaceAfter=4))
                continue

            # Bullet points — escape special XML chars
            safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

            if text.startswith(("•", "-", "–", "*", "·", "○", "▸")):
                bullet_text = safe_text[1:].strip()
                story.append(Paragraph(f"• {bullet_text}", bullet_style))
            else:
                story.append(Paragraph(safe_text, body_style))

        doc_out = SimpleDocTemplate(
            pdf_path,
            pagesize=LETTER,
            rightMargin=0.65 * inch,
            leftMargin=0.65 * inch,
            topMargin=0.5 * inch,
            bottomMargin=0.5 * inch,
        )
        doc_out.build(story)
        return os.path.exists(pdf_path)

    except Exception as e:
        logger.error(f"ReportLab fallback failed: {e}")
        return False


def convert_docx_to_pdf(docx_path: str, output_dir: str) -> Optional[str]:
    """
    Try all PDF conversion methods in priority order.
    Returns the PDF path or None if all methods fail.
    """
    stem = Path(docx_path).stem
    pdf_path = os.path.join(output_dir, f"{stem}.pdf")

    # 1. Try docx2pdf (Word on Mac) - Disabled to prevent Word app popup
    # if _convert_docx_to_pdf_via_docx2pdf(docx_path, pdf_path):
    #     logger.info(f"PDF via docx2pdf: {pdf_path}")
    #     return pdf_path

    # 2. Try LibreOffice
    lo_result = _convert_docx_to_pdf_via_libreoffice(docx_path, output_dir)
    if lo_result:
        # Rename to our expected path if different
        if lo_result != pdf_path and os.path.exists(lo_result):
            os.rename(lo_result, pdf_path)
        logger.info(f"PDF via LibreOffice: {pdf_path}")
        return pdf_path

    # 3. ReportLab clean fallback
    if _convert_docx_to_pdf_reportlab_fallback(docx_path, pdf_path):
        logger.info(f"PDF via ReportLab fallback: {pdf_path}")
        return pdf_path

    return None


# ─── DOCX text replacement (format-preserving) ────────────────────────────────

def _clear_paragraph_text(para) -> None:
    """Remove all runs from a paragraph without touching its style."""
    for run in para.runs:
        run.text = ""


def _set_paragraph_text_preserve_format(para, new_text: str) -> None:
    """
    Replace the text of a paragraph while preserving the formatting
    of the FIRST run (font, size, bold, italic, color).
    Non-empty subsequent runs are cleared.
    """
    if not para.runs:
        para.add_run(new_text)
        return

    # Keep the first run's formatting; set its text to new_text
    first_run = para.runs[0]
    first_run.text = new_text

    # Clear all other runs (avoid duplicate text)
    for run in para.runs[1:]:
        run.text = ""


def extract_docx_mapping(docx_path: str) -> dict:
    """Extract paragraphs into a mapping { '0': 'Text', ... }"""
    doc = Document(docx_path)
    mapping = {}
    for i, para in enumerate(doc.paragraphs):
        mapping[str(i)] = para.text
    return mapping


def generate_optimized_documents_from_mapping(
    optimized_mapping: dict,
    output_dir: str,
    original_file_path: str,
) -> dict[str, Optional[str]]:
    """
    Applies an exact paragraph mapping to the original DOCX file, preserving ALL
    native formatting and layout perfectly, then converts to PDF.
    """
    import uuid
    os.makedirs(output_dir, exist_ok=True)
    uid = uuid.uuid4().hex[:10]
    result = {"docx_path": None, "pdf_path": None}

    try:
        doc = Document(original_file_path)
        for doc_id, text in optimized_mapping.items():
            try:
                idx = int(doc_id)
                para = doc.paragraphs[idx]
                if para.text != text:
                    _set_paragraph_text_preserve_format(para, text)
            except (ValueError, IndexError):
                pass
        
        working_docx_path = os.path.join(output_dir, f"optimized_{uid}.docx")
        doc.save(working_docx_path)
        logger.info(f"Format-preserving DOCX saved: {working_docx_path}")
        result["docx_path"] = working_docx_path

        # Convert DOCX → PDF
        pdf_path = convert_docx_to_pdf(working_docx_path, output_dir)
        result["pdf_path"] = pdf_path

        return result
    except Exception as e:
        logger.error(f"Format-preserving optimization from mapping failed: {e}")
        return result


# ─── PDF→DOCX conversion ──────────────────────────────────────────────────────

def convert_pdf_to_docx(pdf_path: str, output_dir: str) -> Optional[str]:
    """Convert a PDF to DOCX using pdf2docx."""
    try:
        from pdf2docx import Converter
        docx_path = os.path.join(output_dir, Path(pdf_path).stem + "_converted.docx")
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        if os.path.exists(docx_path):
            logger.info(f"PDF→DOCX via pdf2docx: {docx_path}")
            return docx_path
    except Exception as e:
        logger.warning(f"pdf2docx conversion failed: {e}")
    return None


# ─── Main public API ─────────────────────────────────────────────────────────

def generate_template_documents(
    resume_data: dict,
    output_dir: str,
    template_type: str = "modern"
) -> dict[str, Optional[str]]:
    """
    Generate beautiful formatted DOCX and PDF files using predefined templates.
    """
    from app.services.templates import generate_templated_docx
    
    result = {"docx_path": None, "pdf_path": None}
    
    try:
        # Generate DOCX from template builder
        docx_path = generate_templated_docx(resume_data, template_type, output_dir)
        result["docx_path"] = docx_path
        
        # Convert to PDF
        pdf_path = convert_docx_to_pdf(docx_path, output_dir)
        result["pdf_path"] = pdf_path
        
        return result
    except Exception as e:
        logger.error(f"Template generation pipeline failed: {e}", exc_info=True)
        return result

def generate_optimized_documents(
    optimized_text: str,
    output_dir: str,
    original_file_path: Optional[str] = None,
    file_type: Optional[str] = None,
) -> dict[str, Optional[str]]:
    """
    Generate optimized DOCX and PDF files using the scratch generator.
    Note: For format preservation, use generate_optimized_documents_from_mapping.

    Returns:
        {
            "docx_path": str | None,
            "pdf_path":  str | None,
        }
    """
    os.makedirs(output_dir, exist_ok=True)
    uid = uuid.uuid4().hex[:10]
    result = {"docx_path": None, "pdf_path": None}

    working_docx_path: Optional[str] = None

    # ── 1. Always generate clean DOCX from scratch ────────────────────────────
    # (Since format preservation is now handled by the mapping pipeline)

    # ── 2. Fallback: generate clean DOCX from scratch ─────────────────────────
    if not working_docx_path:
        working_docx_path = _generate_clean_docx_from_scratch(optimized_text, output_dir, uid)

    if working_docx_path:
        result["docx_path"] = working_docx_path

        # ── 3. Convert DOCX → PDF ─────────────────────────────────────────────
        pdf_path = convert_docx_to_pdf(working_docx_path, output_dir)
        result["pdf_path"] = pdf_path

    return result


def _generate_clean_docx_from_scratch(text: str, output_dir: str, uid: str) -> Optional[str]:
    """
    Generate a professionally formatted DOCX from the optimized text.
    Used as fallback when we can't preserve original formatting.
    """
    try:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        doc = Document()

        # Set narrow margins
        from docx.shared import Inches
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.65)
            section.right_margin = Inches(0.65)

        SECTION_KEYWORDS = {
            "experience", "education", "skills", "projects", "summary",
            "objective", "certifications", "awards", "work experience",
            "professional experience", "technical skills", "volunteer",
        }

        lines = text.splitlines()
        is_first = True

        for line in lines:
            stripped = line.strip()

            if not stripped:
                # Add small spacing between blocks
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                continue

            lower = stripped.lower().rstrip(":")
            is_section = lower in SECTION_KEYWORDS or (len(stripped) < 60 and stripped.isupper())

            if is_first:
                # Name
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
                p.paragraph_format.space_after = Pt(2)
                is_first = False

                # Underline separator
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(4)
                run2 = p2.add_run("─" * 75)
                run2.font.size = Pt(6)
                run2.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
                continue

            if is_section:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(stripped.upper())
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
                run.font.name = "Calibri"

                # Section divider
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(3)
                run2 = p2.add_run("─" * 75)
                run2.font.size = Pt(5)
                run2.font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)
                continue

            # Bullet points
            if stripped.startswith(("•", "-", "–", "·", "*", "▸", "○")):
                p = doc.add_paragraph(style="List Bullet")
                content = stripped[1:].strip()
                run = p.add_run(content)
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                p.paragraph_format.space_after = Pt(1)
            else:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                p.paragraph_format.space_after = Pt(1)

        docx_path = os.path.join(output_dir, f"optimized_{uid}.docx")
        doc.save(docx_path)
        logger.info(f"Clean DOCX generated from scratch: {docx_path}")
        return docx_path

    except Exception as e:
        logger.error(f"Clean DOCX generation failed: {e}")
        return None
