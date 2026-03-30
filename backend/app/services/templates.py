import os
import uuid
import logging
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Shared helper (used by modern template only)
# ─────────────────────────────────────────────────────────────────────────────

def add_horizontal_line(p, color_hex="D1D5DB", size=5):
    """Adds a visual divider below a paragraph (used by modern/classic/minimal)."""
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("─" * 75)
    run.font.size = Pt(size)
    try:
        r = int(color_hex[0:2], 16)
        g = int(color_hex[2:4], 16)
        b = int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    except Exception:
        run.font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)


# ─────────────────────────────────────────────────────────────────────────────
# MANI template — constants, helpers, skill categoriser
# ─────────────────────────────────────────────────────────────────────────────

# Colors  — reverse-engineered from original .docx
_M_SKILL_CATEGORY_COLOR = RGBColor(0x21, 0x5e, 0x99)  # #215E99 - medium blue
_M_DARK_BLUE  = RGBColor(0x1f, 0x4e, 0x79)   # section headers / name
_M_LINK_COLOR = RGBColor(0x21, 0x5e, 0x99)   # hyperlinks  (#215E99, matches original)
_M_BLACK      = RGBColor(0x00, 0x00, 0x00)
_M_HARDCODED_NAME = "MANIDHEER NALLA"

# Border values reverse-engineered from original .docx XML
_M_BORDER_COLOR = "AAAAAA"   # grey rule matching original
_M_BORDER_SZ    = "10"       # w:sz=10 → 1.25 pt
_M_BORDER_SPACE = "3"

# Typography
_M_FONT        = "Calibri"
_M_NAME_PT     = 18
_M_CONTACT_PT  = 12    # contact section — matches original exactly
_M_HEADER_PT   = 12
_M_BODY_PT     = 10.5
_M_BULLET_PT   = 10
_M_SKILL_CAT_PT = 10.5  # category label in skills table
_M_SKILL_PT    = 10

# Layout
_M_LEFT_M  = 0.75
_M_RIGHT_M = 0.75
_M_TEXT_W  = 8.5 - _M_LEFT_M - _M_RIGHT_M   # 7.0 in — right-tab anchor

# ── Hardcoded contact details ─────────────────────────────────────────────────
_CONTACT = {
    "email":     "manidheerft@gmail.com",
    "phone":     "682-407-5087",
    "location":  "Dallas, TX 75201",
    "linkedin":  "https://www.linkedin.com/in/manireddynalla",
    "github":    "https://github.com/manidheer",
    "portfolio": "https://portfolio.manidheer.me/",
}

# ── Skills that should be rendered BOLD in the skills table ──────────────────
# Mirrors the selective bolding in the original resume.
_M_BOLD_SKILLS = {
    # AI / ML
    "pytorch", "llms", "llm", "generative ai", "rag", "langchain", "llamaindex",
    "ai agents", "langgraph", "crewai", "prompt engineering", "embeddings",
    "semantic search", "mlflow", "tensorflow", "dspy", "pinecone", "chroma", "faiss",
    "huggingface", "transformers",
    # Backend
    "c#", ".net", "asp.net", "fastapi", "rest apis", "rest api", "graphql",
    "microservices", "grpc",
    # Frontend
    "react", "javascript", "typescript", "blazor", "next.js", "nextjs",
    # DB
    "postgresql", "pgvector", "redis", "mongodb", "elasticsearch",
    # Cloud
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
    # Tools
    "linux / bash", "linux", "bash", "git", "pytest",
}

# Longest-first to avoid short terms preempting longer matches (e.g., llm vs llms).
_M_BOLD_SKILLS_SORTED = sorted(_M_BOLD_SKILLS, key=len, reverse=True)
_M_BOLD_SKILL_REGEX = re.compile(
    r"(?i)(?<![A-Za-z0-9])(" + "|".join(re.escape(s) for s in _M_BOLD_SKILLS_SORTED) + r")(?![A-Za-z0-9])"
)

# ── Skill-category keyword map ────────────────────────────────────────────────
_SKILL_CATEGORIES = {
    "AI / Machine Learning": [
        "python", "pytorch", "tensorflow", "llms", "llm", "generative ai",
        "openai", "anthropic", "gemini", "rag", "langchain", "llamaindex",
        "ai agents", "langgraph", "crewai", "prompt engineering", "dspy",
        "pinecone", "chroma", "faiss", "vector database", "vector databases",
        "embeddings", "semantic search", "pandas", "numpy", "scikit-learn",
        "mlflow", "huggingface", "transformers", "spacy", "nltk",
    ],
    "Backend Development": [
        "c#", ".net", "asp.net", "fastapi", "django", "node.js", "nodejs",
        "rest apis", "rest api", "graphql", "microservices", "jwt", "oauth2",
        "flask", "express", "spring", "java", "go", "ruby", "rails", "grpc",
        "celery", "rabbitmq",
    ],
    "Frontend Development": [
        "react", "javascript", "next.js", "nextjs", "typescript", "blazor",
        "webassembly", "tailwind css", "tailwind", "chart.js", "vue",
        "angular", "svelte", "html", "css", "sass", "webpack", "vite",
    ],
    "Databases & Data Systems": [
        "postgresql", "pgvector", "sql server", "mongodb", "redis",
        "entity framework", "sqlalchemy", "mysql", "sqlite", "cassandra",
        "dynamodb", "elasticsearch", "snowflake", "bigquery", "dbt",
        "airflow", "kafka", "spark",
    ],
    "Cloud & DevOps": [
        "aws", "azure", "gcp", "docker", "kubernetes", "ci/cd",
        "github actions", "terraform", "nginx", "jenkins", "gitlab",
        "circleci", "ansible", "heroku", "vercel", "netlify",
    ],
    "Developer Tools": [
        "git", "linux", "bash", "linux / bash", "poetry", "pytest",
        "langsmith", "jupyter", "jupyter notebook", "postman", "vs code",
        "intellij", "jira", "confluence", "figma", "datadog", "sentry",
    ],
}


def _m_categorize_skills(flat_skills: list) -> dict:
    """
    Re-buckets the optimizer's flat ['Python','React',...] list into the
    category dict used by the 2-column skills table.

    Two-pass strategy to prevent short-keyword false matches
    (e.g. 'java' matching 'JavaScript', 'go' matching 'MongoDB',
    'git' matching 'github actions'):
      Pass 1 — exact match across ALL categories first.
      Pass 2 — substring match, only if no exact match was found.
    """
    buckets = {cat: [] for cat in _SKILL_CATEGORIES}
    buckets["Other"] = []
    for skill in flat_skills:
        skill_lower = skill.strip().lower()
        placed = False

        # Pass 1: exact match — checks every category before falling back
        for cat, keywords in _SKILL_CATEGORIES.items():
            if skill_lower in keywords:
                buckets[cat].append(skill.strip())
                placed = True
                break

        # Pass 2: substring match (only if no exact match found)
        if not placed:
            for cat, keywords in _SKILL_CATEGORIES.items():
                if any(kw in skill_lower or skill_lower in kw for kw in keywords):
                    buckets[cat].append(skill.strip())
                    placed = True
                    break

        if not placed:
            buckets["Other"].append(skill.strip())

    return {cat: skills for cat, skills in buckets.items() if skills}


def _m_sp(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)


def _m_run(para, text, size=_M_BODY_PT, bold=False, color=None):
    """Add a plain run with consistent font settings."""
    r = para.add_run(str(text))
    r.font.size      = Pt(size)
    r.font.name      = _M_FONT
    r.font.color.rgb = color or _M_BLACK
    r.bold           = bold
    return r


def _m_section_header(doc, title, space_before=0.5):
    """
    Single paragraph with top+bottom grey borders — matches the original .docx
    exactly. Using ONE paragraph (not three) eliminates the large gap that the
    old 'empty divider para → header → empty divider para' approach caused.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _m_sp(p, before=space_before, after=1)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    _M_BORDER_SZ)
        el.set(qn('w:space'), _M_BORDER_SPACE)
        el.set(qn('w:color'), _M_BORDER_COLOR)
        pBdr.append(el)
    pPr.append(pBdr)

    _m_run(p, title.upper(), size=_M_HEADER_PT, bold=True, color=_M_DARK_BLUE)
    return p


def _m_add_hyperlink(para, text, url, bold=False, size=_M_BODY_PT):
    """Insert a real clickable hyperlink into an existing paragraph."""
    part = para.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    rPr_el = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr_el.append(rStyle)

    # Override font properties to match original exactly
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), _M_FONT)
    rFonts.set(qn('w:hAnsi'), _M_FONT)
    rPr_el.append(rFonts)

    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), str(int(size * 2)))
    rPr_el.append(sz_el)

    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), f'{_M_LINK_COLOR[0]:02X}{_M_LINK_COLOR[1]:02X}{_M_LINK_COLOR[2]:02X}')
    rPr_el.append(color_el)

    # Force visible underline on links to match the requested heading style.
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr_el.append(u_el)

    if bold:
        b_el = OxmlElement('w:b')
        rPr_el.append(b_el)

    r_el = OxmlElement('w:r')
    r_el.append(rPr_el)
    t_el = OxmlElement('w:t')
    t_el.text = text
    r_el.append(t_el)
    hyperlink.append(r_el)
    para._p.append(hyperlink)


def _m_job_line(doc, title, company, duration):
    """Bold 'Title — Company' with right-tab-aligned date."""
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(_M_TEXT_W), WD_TAB_ALIGNMENT.RIGHT)
    _m_sp(p, before=1, after=1)
    _m_run(p, f"{title} — {company}", size=_M_BODY_PT, bold=True)
    if duration:
        p.add_run("\t")
        _m_run(p, duration, size=_M_BODY_PT, bold=True)


def _m_run_with_bold(para, text: str, size=_M_BODY_PT, color=None, auto_bold_skills=True):
    """
    Add text to `para` parsing **bold** markdown markers into bold runs.
    Plain segments get regular weight; segments inside **...** get bold.
    """
    parts = re.split(r'\*\*([^*]+)\*\*', text)

    def _append_runs(segment: str, force_bold: bool = False):
        if not segment:
            return
        if force_bold or not auto_bold_skills:
            r = para.add_run(segment)
            r.font.size = Pt(size)
            r.font.name = _M_FONT
            r.font.color.rgb = color or _M_BLACK
            r.bold = force_bold
            return

        last = 0
        for m in _M_BOLD_SKILL_REGEX.finditer(segment):
            start, end = m.span()
            if start > last:
                normal_run = para.add_run(segment[last:start])
                normal_run.font.size = Pt(size)
                normal_run.font.name = _M_FONT
                normal_run.font.color.rgb = color or _M_BLACK
                normal_run.bold = False

            bold_run = para.add_run(segment[start:end])
            bold_run.font.size = Pt(size)
            bold_run.font.name = _M_FONT
            bold_run.font.color.rgb = color or _M_BLACK
            bold_run.bold = True
            last = end

        if last < len(segment):
            tail_run = para.add_run(segment[last:])
            tail_run.font.size = Pt(size)
            tail_run.font.name = _M_FONT
            tail_run.font.color.rgb = color or _M_BLACK
            tail_run.bold = False

    for i, part in enumerate(parts):
        if not part:
            continue
        # odd indices are markdown **...** groups and should remain bold as-is.
        _append_runs(part, force_bold=(i % 2 == 1))


def _m_bullets(doc, raw):
    """
    Renders bullet items. Accepts either:
      - a list of strings  (standard optimizer output)
      - a single \\n•-joined string  (legacy format)
    """
    if isinstance(raw, str):
        lines = [l.lstrip('•').strip() for l in raw.replace('\r', '').split('\n')]
        raw = [l for l in lines if l]
    for item in raw:
        item = item.lstrip('•').strip()
        if not item:
            continue
        try:
            pb = doc.add_paragraph(style='List Bullet')
        except KeyError:
            pb = doc.add_paragraph(style='List Paragraph')
            pb.paragraph_format.left_indent = Inches(0.25)
        _m_sp(pb, before=0, after=1)
        _m_run_with_bold(pb, item, size=_M_BULLET_PT)


def _m_render_skills_with_bold(para, skills_list):
    """
    Renders a list of skills into `para` with selective bold on key technologies.
    Matches the original resume's inline run-level bold formatting.
    """
    for i, skill in enumerate(skills_list):
        skill_stripped = skill.strip()
        is_bold = skill_stripped.lower() in _M_BOLD_SKILLS
        text = skill_stripped + (", " if i < len(skills_list) - 1 else "")
        r = para.add_run(text)
        r.bold = is_bold
        r.font.size = Pt(_M_SKILL_PT)
        r.font.name = _M_FONT
        r.font.color.rgb = _M_BLACK


# ─────────────────────────────────────────────────────────────────────────────
# MANI template — main function
# ─────────────────────────────────────────────────────────────────────────────

def apply_mani_template(doc: Document, data: dict):
    """
    Pixel-perfect Manidheer Nalla resume — matches the original .docx exactly.

    Changes from previous version:
    1. Contact section: size 12pt, order Location | Phone | Email (matches original)
    2. Hyperlinks: bold, size 12pt, color #215E99 (matches original)
    3. Skills table: category labels 10.5pt bold, skills 10.0pt with selective bold
    4. Spacer paragraph (sb=2) after skills table before EXPERIENCE header
    5. EXPERIENCE section header uses sb=0 (spacer provides the gap)
    6. Education: single-line "Degree | Institution — City  | Year" format
    7. Projects: name bold 10.5pt, description plain 10.0pt
    """

    # ── Margins ───────────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin   = Inches(_M_LEFT_M)
        sec.right_margin  = Inches(_M_RIGHT_M)

    # Name is intentionally hardcoded for the Mani template.
    name         = _M_HARDCODED_NAME
    summary_body = (data.get('summary') or '').strip()

    # ── NAME ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _m_sp(p, before=0, after=1)
    _m_run(p, name.upper(), size=_M_NAME_PT, bold=True, color=_M_DARK_BLUE)

    # ── CONTACT ───────────────────────────────────────────────────────────────
    # Line 1: Location | Phone | Email  (size 12pt, matches original exactly)
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _m_sp(p1, before=0, after=1)
    _m_run(p1,
           f"{_CONTACT['location']} | {_CONTACT['phone']} | ",
           size=_M_CONTACT_PT)
    _m_add_hyperlink(p1, _CONTACT['email'], f"mailto:{_CONTACT['email']}", size=_M_CONTACT_PT)

    # Line 2: Bold clickable hyperlinks — LinkedIn | Portfolio | GitHub
    # Bold + size 12 + color #215E99 matches the original .docx runs exactly
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _m_sp(p2, before=0, after=1)
    _m_add_hyperlink(p2, "LinkedIn",  _CONTACT['linkedin'],  bold=True, size=_M_CONTACT_PT)
    _m_run(p2, "  |  ", size=_M_CONTACT_PT, bold=True, color=_M_LINK_COLOR)
    _m_add_hyperlink(p2, "Portfolio", _CONTACT['portfolio'], bold=True, size=_M_CONTACT_PT)
    _m_run(p2, "  |  ", size=_M_CONTACT_PT, bold=True, color=_M_LINK_COLOR)
    _m_add_hyperlink(p2, "GitHub",    _CONTACT['github'],    bold=True, size=_M_CONTACT_PT)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    if summary_body:
        _m_section_header(doc, 'SUMMARY')
        p = doc.add_paragraph()
        _m_sp(p, before=0, after=1)
        # Keep summary plain text (no keyword bolding).
        summary_plain = re.sub(r'\*\*([^*]+)\*\*', r'\1', summary_body)
        _m_run(p, summary_plain, size=_M_BODY_PT)

    # ── TECHNICAL SKILLS ──────────────────────────────────────────────────────
    skills = data.get('skills')
    if skills:
        _m_section_header(doc, 'TECHNICAL SKILLS')

        # Normalise to category dict regardless of input format
        if isinstance(skills, list) and skills and ':' not in str(skills[0]):
            # Flat list from optimizer → re-bucket into categories
            skill_dict = _m_categorize_skills(skills)
        elif isinstance(skills, dict):
            skill_dict = skills
        else:
            # Already 'Category: items' strings
            skill_dict = {}
            for s in (skills if isinstance(skills, list) else [str(skills)]):
                parts = str(s).split(':', 1)
                skill_dict[parts[0].strip()] = parts[1].strip() if len(parts) > 1 else s

        # Build list of (category_label, [skill1, skill2, ...]) tuples
        cat_skills = []
        for cat, vals in skill_dict.items():
            if isinstance(vals, list):
                items = vals
            else:
                # vals is a comma-string like "Python, FastAPI, ..."
                items = [s.strip() for s in str(vals).split(',') if s.strip()]
            cat_skills.append((cat, items))

        mid         = (len(cat_skills) + 1) // 2
        left_cats   = cat_skills[:mid]
        right_cats  = cat_skills[mid:]

        # Borderless 2-column table — matches original "Normal Table" style
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        tblPr = tbl._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            tblBorders.append(el)
        tblPr.append(tblBorders)

        def _fill_cell(cell, cats):
            """
            Each category = two paragraphs:
              p0: bold category label  (10.5 pt)  — matches original exactly
              p1: skill items with selective bold (10.0 pt)
            """
            first = True
            for cat_label, items in cats:
                # Category header paragraph
                cp = cell.paragraphs[0] if first else cell.add_paragraph()
                first = False
                _m_sp(cp, before=0, after=1)
                rc = cp.add_run(cat_label + ':')
                rc.bold = True
                rc.font.size = Pt(_M_SKILL_CAT_PT)
                rc.font.name = _M_FONT
                rc.font.color.rgb = _M_SKILL_CATEGORY_COLOR

                # Skills paragraph
                if items:
                    sp = cell.add_paragraph()
                    _m_sp(sp, before=0, after=1)
                    _m_render_skills_with_bold(sp, items)

        _fill_cell(tbl.rows[0].cells[0], left_cats)
        _fill_cell(tbl.rows[0].cells[1], right_cats)

        # Spacer paragraph after table — matches original (sb=0.5, sa=0)
        spacer = doc.add_paragraph()
        _m_sp(spacer, before=0.5, after=0)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    if data.get('experience'):
        _m_section_header(doc, 'EXPERIENCE')
        for exp in data['experience']:
            _m_job_line(doc,
                        exp.get('title', ''),
                        exp.get('company', ''),
                        exp.get('duration', ''))
            _m_bullets(doc, exp.get('bullets', []))

    # ── AI PROJECTS ───────────────────────────────────────────────────────────
    if data.get('projects'):
        _m_section_header(doc, 'AI PROJECTS')
        for proj in data['projects']:
            # Project name — bold 10.5pt
            p_name = doc.add_paragraph()
            _m_sp(p_name, before=1, after=0)
            _m_run(p_name, proj.get('name', ''), size=_M_BODY_PT, bold=True)

            # Project description — plain paragraph (matches original .docx layout)
            desc = proj.get('description', '')
            if not desc:
                # Legacy fallback: join bullets into a single paragraph
                bullets = proj.get('bullets', [])
                if isinstance(bullets, list):
                    desc = ' '.join(b.lstrip('•').strip() for b in bullets if b.strip())
                elif isinstance(bullets, str):
                    desc = bullets.lstrip('•').strip()
            if desc:
                p_desc = doc.add_paragraph()
                _m_sp(p_desc, before=0, after=1)
                _m_run_with_bold(p_desc, desc.strip(), size=_M_BULLET_PT)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    # Single-line format: "Degree Field | Institution — City, State  | Year"
    # Matches the original .docx exactly (all inline, no tab-stop for year)
    if data.get('education'):
        _m_section_header(doc, 'EDUCATION')
        for edu in data['education']:
            degree      = f"{edu.get('degree', '')} {edu.get('field', '')}".strip()
            institution = edu.get('institution', '')
            year        = edu.get('year', '')

            p = doc.add_paragraph()
            _m_sp(p, before=0, after=1)

            # Build single-line string matching original layout
            parts = [degree]
            if institution:
                parts.append(institution)
            line = ' | '.join(parts)
            if year:
                line += f'  | {year}'

            _m_run(p, line, size=_M_BODY_PT, bold=True)


# ─────────────────────────────────────────────────────────────────────────────
# Classic / Modern templates
# ─────────────────────────────────────────────────────────────────────────────

def _normalize_bullet_list(raw):
    """Normalize bullets from list or newline text into a clean list of strings."""
    if isinstance(raw, str):
        lines = [l.lstrip('•').strip() for l in raw.replace('\r', '').split('\n')]
        return [l for l in lines if l]
    if isinstance(raw, list):
        return [str(i).lstrip('•').strip() for i in raw if str(i).strip()]
    return []


def _normalize_skill_lines(skills):
    """Return skills as a list of 'Category: a, b' lines for non-mani templates."""
    if not skills:
        return []

    if isinstance(skills, dict):
        lines = []
        for cat, vals in skills.items():
            if isinstance(vals, list):
                lines.append(f"{cat}: {', '.join(v.strip() for v in vals if str(v).strip())}")
            else:
                lines.append(f"{cat}: {str(vals).strip()}")
        return [l for l in lines if l.strip()]

    if isinstance(skills, list):
        cleaned = [str(s).strip() for s in skills if str(s).strip()]
        if not cleaned:
            return []
        if any(':' in item for item in cleaned):
            return cleaned

        # Keep non-mani templates independent from Mani skill categories.
        chunk_size = 8
        return [", ".join(cleaned[i:i + chunk_size]) for i in range(0, len(cleaned), chunk_size)]

    return [str(skills).strip()] if str(skills).strip() else []


def _extract_name_contact(data: dict):
    """Build display name/contact for classic and modern templates."""
    name = (data.get('name') or data.get('full_name') or '').strip()
    if not name:
        name = "CANDIDATE NAME"

    contact_info = (data.get('contact_info') or '').strip()
    if contact_info:
        return name, contact_info

    contact = data.get('contact') if isinstance(data.get('contact'), dict) else {}
    email = str(contact.get('email') or data.get('email') or '').strip()
    phone = str(contact.get('phone') or data.get('phone') or '').strip()
    location = str(contact.get('location') or data.get('location') or '').strip()

    parts = [value for value in [location, phone, email] if value]
    return name, " | ".join(parts)


def _extract_contact_links(data: dict):
    """Return non-empty contact links for non-mani templates."""
    contact = data.get('contact') if isinstance(data.get('contact'), dict) else {}
    linkedin = str(contact.get('linkedin') or data.get('linkedin') or '').strip()
    portfolio = str(contact.get('portfolio') or data.get('portfolio') or '').strip()
    github = str(contact.get('github') or data.get('github') or '').strip()
    return {
        'linkedin': linkedin,
        'portfolio': portfolio,
        'github': github,
    }


def _classic_section_header(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _m_sp(p, before=7, after=1)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.font.color.rgb = _M_BLACK

    divider = doc.add_paragraph()
    _m_sp(divider, before=0, after=2)
    line = divider.add_run("_" * 90)
    line.font.size = Pt(8)
    line.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def apply_classic_template(doc: Document, data: dict):
    """Classic, conservative serif resume layout."""
    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    name, contact_line = _extract_name_contact(data)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _m_sp(p_name, before=0, after=2)
    r_name = p_name.add_run(name.upper())
    r_name.bold = True
    r_name.font.size = Pt(18)
    r_name.font.name = "Times New Roman"

    if contact_line:
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _m_sp(p_contact, before=0, after=6)
        r_contact = p_contact.add_run(contact_line)
        r_contact.font.size = Pt(10.5)
        r_contact.font.name = "Times New Roman"

    summary = (data.get('summary') or '').strip()
    if summary:
        _classic_section_header(doc, "Summary")
        p = doc.add_paragraph(summary)
        _m_sp(p, before=0, after=2)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(10.5)

    skill_lines = _normalize_skill_lines(data.get('skills'))
    if skill_lines:
        _classic_section_header(doc, "Technical Skills")
        for line in skill_lines:
            p = doc.add_paragraph(line)
            _m_sp(p, before=0, after=1)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(10.5)

    if data.get('experience'):
        _classic_section_header(doc, "Experience")
        for exp in data['experience']:
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.8), WD_TAB_ALIGNMENT.RIGHT)
            _m_sp(p, before=3, after=1)
            left = f"{exp.get('title', '')} | {exp.get('company', '')}".strip(' |')
            r1 = p.add_run(left)
            r1.bold = True
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(11)
            if exp.get('duration'):
                p.add_run("\t")
                r2 = p.add_run(exp.get('duration', ''))
                r2.bold = True
                r2.font.name = "Times New Roman"
                r2.font.size = Pt(10.5)

            for b in _normalize_bullet_list(exp.get('bullets', [])):
                pb = doc.add_paragraph(style='List Bullet')
                _m_sp(pb, before=0, after=1)
                rb = pb.add_run(b)
                rb.font.name = "Times New Roman"
                rb.font.size = Pt(10.5)

    if data.get('projects'):
        _classic_section_header(doc, "Projects")
        for proj in data['projects']:
            p = doc.add_paragraph()
            _m_sp(p, before=3, after=0)
            rn = p.add_run(proj.get('name', ''))
            rn.bold = True
            rn.font.name = "Times New Roman"
            rn.font.size = Pt(10.5)

            desc = (proj.get('description') or '').strip()
            if not desc:
                desc = ' '.join(_normalize_bullet_list(proj.get('bullets', [])))
            if desc:
                pd = doc.add_paragraph(desc)
                _m_sp(pd, before=0, after=1)
                for r in pd.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10.5)

    if data.get('education'):
        _classic_section_header(doc, "Education")
        for edu in data['education']:
            degree = f"{edu.get('degree', '')} {edu.get('field', '')}".strip()
            institution = edu.get('institution', '')
            year = edu.get('year', '')
            line = ' | '.join([x for x in [degree, institution, year] if x])
            p = doc.add_paragraph(line)
            _m_sp(p, before=1, after=0)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(10.5)
                r.bold = True


def _modern_section_header(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _m_sp(p, before=7, after=0)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = _M_DARK_BLUE
    add_horizontal_line(p, color_hex="C7D2FE", size=5)


def apply_modern_template(doc: Document, data: dict):
    """Modern clean layout with blue accents and compact spacing."""
    for sec in doc.sections:
        sec.top_margin = Inches(0.55)
        sec.bottom_margin = Inches(0.55)
        sec.left_margin = Inches(0.65)
        sec.right_margin = Inches(0.65)

    name, contact_line = _extract_name_contact(data)
    links = _extract_contact_links(data)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _m_sp(p_name, before=0, after=2)
    rn = p_name.add_run(name.upper())
    rn.bold = True
    rn.font.size = Pt(20)
    rn.font.name = _M_FONT
    rn.font.color.rgb = _M_DARK_BLUE

    if contact_line or any(links.values()):
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _m_sp(p_contact, before=0, after=5)

        if contact_line:
            _m_run(p_contact, contact_line, size=10)

        link_parts = []
        if links['linkedin']:
            link_parts.append(("LinkedIn", links['linkedin']))
        if links['portfolio']:
            link_parts.append(("Portfolio", links['portfolio']))
        if links['github']:
            link_parts.append(("GitHub", links['github']))

        if link_parts:
            if contact_line:
                _m_run(p_contact, "  |  ", size=10, color=_M_LINK_COLOR)
            for idx, (label, url) in enumerate(link_parts):
                _m_add_hyperlink(p_contact, label, url, size=10)
                if idx < len(link_parts) - 1:
                    _m_run(p_contact, "  |  ", size=10, color=_M_LINK_COLOR)

    summary = (data.get('summary') or '').strip()
    if summary:
        _modern_section_header(doc, "Summary")
        p = doc.add_paragraph()
        _m_sp(p, before=1, after=2)
        _m_run(p, re.sub(r'\*\*([^*]+)\*\*', r'\1', summary), size=10.5)

    skill_lines = _normalize_skill_lines(data.get('skills'))
    if skill_lines:
        _modern_section_header(doc, "Technical Skills")
        for line in skill_lines:
            p = doc.add_paragraph()
            _m_sp(p, before=1, after=1)
            _m_run_with_bold(p, line, size=10, auto_bold_skills=False)

    if data.get('experience'):
        _modern_section_header(doc, "Experience")
        for exp in data['experience']:
            _m_job_line(doc, exp.get('title', ''), exp.get('company', ''), exp.get('duration', ''))
            _m_bullets(doc, exp.get('bullets', []))

    if data.get('projects'):
        _modern_section_header(doc, "Projects")
        for proj in data['projects']:
            p = doc.add_paragraph()
            _m_sp(p, before=3, after=0)
            _m_run(p, proj.get('name', ''), size=10.5, bold=True, color=_M_DARK_BLUE)

            desc = (proj.get('description') or '').strip()
            if not desc:
                desc = ' '.join(_normalize_bullet_list(proj.get('bullets', [])))
            if desc:
                pd = doc.add_paragraph()
                _m_sp(pd, before=0, after=1)
                _m_run_with_bold(pd, desc, size=10)

    if data.get('education'):
        _modern_section_header(doc, "Education")
        for edu in data['education']:
            degree = f"{edu.get('degree', '')} {edu.get('field', '')}".strip()
            institution = edu.get('institution', '')
            year = edu.get('year', '')
            line = ' | '.join([x for x in [degree, institution, year] if x])
            p = doc.add_paragraph()
            _m_sp(p, before=1, after=0)
            _m_run(p, line, size=10.5, bold=True)


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

def generate_templated_docx(data: dict, template_type: str, output_dir: str) -> str:
    """Takes structured resume JSON data and creates a downloadable DOCX."""
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    template_key = (template_type or "mani").strip().lower()
    if template_key == "mani":
        apply_mani_template(doc, data)
    elif template_key == "modern":
        apply_modern_template(doc, data)
    elif template_key == "classic":
        apply_classic_template(doc, data)
    else:
        logger.warning("Unknown template_type '%s'. Falling back to mani.", template_type)
        apply_mani_template(doc, data)

    uid = uuid.uuid4().hex[:10]
    file_path = os.path.join(output_dir, f"resume_{template_type}_{uid}.docx")
    doc.save(file_path)
    return file_path
